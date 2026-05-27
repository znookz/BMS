"""Script to generate BMS SRS document in Thai for client."""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── Page margins ───
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Styles ───
BRAND_BLUE = RGBColor(0x1E, 0x3A, 0x8A)
BRAND_LIGHT = RGBColor(0xDB, 0xEA, 0xFE)
GRAY = RGBColor(0x6B, 0x72, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x11, 0x18, 0x27)

FONT = "TH Sarabun New"

def set_font(run, size=14, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_font(run, 20, bold=True, color=BRAND_BLUE)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1E3A8A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run, 16, bold=True, color=BRAND_BLUE)
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, 14, bold=True, color=BLACK)
    return p

def body(text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    set_font(run, 14)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run, 14)
    return p

def note_box(text):
    """Yellow note paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("⚠ หมายเหตุ: " + text)
    set_font(run, 13, italic=True, color=RGBColor(0x92, 0x40, 0x09))
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=13, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size, bold=bold, color=color)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, "1E3A8A")
        set_cell_text(cell, h, bold=True, size=13, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg = "DBEAFE" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, bg)
            set_cell_text(cell, val, size=13)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


# ═══════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("เอกสารข้อกำหนดความต้องการซอฟต์แวร์")
set_font(run, 28, bold=True, color=BRAND_BLUE)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Software Requirements Specification (SRS)")
set_font(run, 18, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("ระบบจัดการรถบัส (Bus Management System — BMS)")
set_font(run, 22, bold=True, color=BLACK)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("เวอร์ชัน 1.0")
set_font(run, 14, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"วันที่: {datetime.date.today().strftime('%d %B %Y')}")
set_font(run, 14, color=GRAY)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Status box (simple table)
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
meta = [
    ("จัดทำโดย", "ทีมพัฒนาระบบ"),
    ("จัดทำให้กับ", "ลูกค้า / ผู้ว่าจ้าง"),
    ("โครงการ", "ระบบจัดการรถบัส (BMS)"),
    ("สถานะ", "Draft — รอการอนุมัติ"),
]
for i, (k, v) in enumerate(meta):
    bg = "1E3A8A" if i == 0 else ("DBEAFE" if i % 2 == 0 else "FFFFFF")
    fc = WHITE if i == 0 else BLACK
    shade_cell(tbl.rows[i].cells[0], bg if i == 0 else "EFF6FF")
    shade_cell(tbl.rows[i].cells[1], "FFFFFF")
    set_cell_text(tbl.rows[i].cells[0], k, bold=True, color=BRAND_BLUE, size=13)
    set_cell_text(tbl.rows[i].cells[1], v, size=13)
for row in tbl.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(9)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════
# 1. บทนำ
# ═══════════════════════════════════════════════════════════
heading1("1. บทนำ")

heading2("1.1 วัตถุประสงค์ของเอกสาร")
body("เอกสารนี้อธิบายความต้องการเชิงฟังก์ชัน และความต้องการที่ไม่ใช่เชิงฟังก์ชันของ ระบบจัดการรถบัส (BMS) "
     "รวมถึงสิ่งที่ลูกค้าต้องเตรียมและมอบให้กับทีมพัฒนา เพื่อให้โครงการดำเนินการได้อย่างราบรื่นและตรงตามความคาดหวัง")

heading2("1.2 ขอบเขตของระบบ")
body("ระบบ BMS เป็นเว็บแอปพลิเคชันสำหรับจัดการการปฏิบัติงานของรถบัส ครอบคลุมโรงงาน 3 แห่ง ได้แก่ "
     "บางปะอิน · อยุธยา · ลพบุรี มีรถบัสรวม 68 คัน แบ่งเป็นรถแอร์ และรถพัดลม ใช้งานผ่านเว็บเบราว์เซอร์บนคอมพิวเตอร์เป็นหลัก "
     "รองรับการใช้งานบนมือถือ (Mobile Responsive)")

heading2("1.3 ผู้ใช้งานระบบ")
add_table(
    ["บทบาท", "สิทธิ์การเข้าถึง", "จำนวนผู้ใช้ (โดยประมาณ)"],
    [
        ("Admin", "เข้าถึงได้ทุกโมดูล รวมถึง User Management และ Audit Log", "—"),
        ("Operation", "เข้าถึงได้ทุกโมดูล ยกเว้น User Management", "—"),
    ],
    col_widths=[3.5, 9, 4],
)
note_box("ลูกค้าต้องระบุจำนวนผู้ใช้งานจริงในแต่ละบทบาท และอีเมลของ Admin คนแรก")


# ═══════════════════════════════════════════════════════════
# 2. สิ่งที่ลูกค้าต้องเตรียม (สำคัญที่สุด)
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1("2. สิ่งที่ลูกค้าต้องเตรียมให้ทีมพัฒนา")

body("ส่วนนี้รวบรวมข้อมูลและการดำเนินการทั้งหมดที่ทีมลูกค้าต้องจัดเตรียม "
     "หากข้อมูลส่วนใดขาดหายหรือล่าช้า อาจส่งผลให้ Timeline ของโครงการล่าช้าตามไปด้วย")

heading2("2.1 ข้อมูลธุรกิจและองค์กร")
add_table(
    ["รายการ", "รายละเอียด", "ผู้รับผิดชอบ (ลูกค้า)", "กำหนดส่ง"],
    [
        ("ชื่อองค์กร / บริษัท", "ชื่อเต็มภาษาไทยและอังกฤษ สำหรับแสดงในระบบและเอกสาร", "", "ก่อนเริ่ม Phase 1"),
        ("โลโก้บริษัท", "ไฟล์ PNG/SVG พื้นหลังโปร่งใส ขนาด ≥ 512×512 px", "", "ก่อนเริ่ม Phase 1"),
        ("รายชื่อโรงงาน/สาขา", "ชื่อ-ที่อยู่ของทั้ง 3 โรงงาน", "", "ก่อนเริ่ม Phase 1"),
        ("รายชื่อบริษัทรถบัส", "7 บริษัท ชื่อย่อ + ชื่อเต็ม สำหรับระบบติดตามน้ำมัน", "", "ก่อนเริ่ม Phase 4"),
        ("จำนวนรถและทะเบียนรถ", "ข้อมูลรถ 68 คัน (ทะเบียน ยี่ห้อ รุ่น ประเภท สังกัดโรงงาน)", "", "ก่อนเริ่ม Phase 3"),
        ("ผังองค์กร / ผู้ใช้งานระบบ", "รายชื่อ + อีเมล + บทบาท ของผู้ใช้งานทุกคน", "", "ก่อนเริ่ม Phase 1"),
    ],
    col_widths=[4.5, 6.5, 3.5, 3],
)

heading2("2.2 ข้อมูลสำหรับนำเข้าระบบ (Master Data)")
body("ข้อมูลเหล่านี้ต้องส่งในรูปแบบ Excel (.xlsx) ตามแบบฟอร์มที่ทีมพัฒนาจะจัดส่งให้")
add_table(
    ["โมดูล", "ข้อมูลที่ต้องการ", "รูปแบบ"],
    [
        ("พนักงานขับรถ (Driver)", "ชื่อ-นามสกุล, รหัสพนักงาน, เบอร์โทร, ใบขับขี่ (ประเภท/วันหมดอายุ), สังกัดโรงงาน", "Excel"),
        ("รถบัส (Bus)", "ทะเบียน, ยี่ห้อ, รุ่น, ปีที่ผลิต, ประเภท (แอร์/พัดลม), สถานะ, วันครบประกัน", "Excel"),
        ("ประวัติซ่อมบำรุง", "ข้อมูลย้อนหลังที่ต้องการ migrate เข้าระบบ (ถ้ามี)", "Excel"),
        ("หลักสูตรอบรม (Training)", "รายชื่อหลักสูตร, ระยะเวลาที่ใช้ได้ (เดือน), ผู้ให้การอบรม", "Excel"),
        ("ข้อมูลน้ำมัน (Fuel/ATG)", "ประวัติการเติมน้ำมันย้อนหลัง (ถ้าต้องการ migrate)", "Excel"),
        ("เส้นทาง/รอบเดินรถ", "รายชื่อเส้นทาง, ระยะทาง, โรงงานต้นทาง-ปลายทาง", "Excel"),
    ],
    col_widths=[4.5, 8, 2.5],
)

heading2("2.3 ข้อมูลทางเทคนิค")
add_table(
    ["รายการ", "รายละเอียด", "หมายเหตุ"],
    [
        ("อีเมลสำหรับ Admin แรก", "ใช้สมัคร Supabase Auth", "ต้องเป็นอีเมลจริง"),
        ("โดเมนเนม (ถ้ามี)", "เช่น bms.yourcompany.com (กรณีต้องการใช้โดเมนของตัวเอง)", "ถ้าไม่มี ใช้โดเมน Vercel"),
        ("ผู้ดูแลระบบ IT (ถ้ามี)", "ชื่อ-เบอร์ติดต่อ IT ที่จะประสานงานด้าน Network / DNS", ""),
        ("ระบบ ATG (น้ำมัน)", "API endpoint, credentials, รูปแบบข้อมูลที่ส่งออกมา", "สำหรับ Phase อนาคต"),
        ("ระบบพฤติกรรมการขับขี่", "API หรือรูปแบบข้อมูล ถ้ามีการ integrate ในอนาคต", "สำหรับ Phase อนาคต"),
    ],
    col_widths=[4.5, 7.5, 3.5],
)

heading2("2.4 ไฟล์และสื่อดิจิทัล")
add_table(
    ["รายการ", "รายละเอียด", "ใช้ใน Module"],
    [
        ("โลโก้บริษัท", "PNG/SVG พื้นโปร่งใส ≥ 512 px", "ทุกหน้า"),
        ("ภาพรถบัส (ถ้ามี)", "รูปตัวอย่างรถแต่ละประเภท สำหรับ UI", "Bus Module"),
        ("เอกสารประกันภัย (ตัวอย่าง)", "เพื่อออกแบบฟิลด์ข้อมูลให้ตรงกับเอกสารจริง", "License & Insurance"),
        ("แบบฟอร์มตรวจสภาพก่อนออกรถ", "แบบฟอร์มกระดาษที่ใช้อยู่ปัจจุบัน", "Safety Measurement"),
        ("แบบฟอร์มบันทึกอุบัติเหตุ", "แบบฟอร์มปัจจุบัน", "Maintenance"),
    ],
    col_widths=[4.5, 7, 4],
)

heading2("2.5 กฎเกณฑ์ทางธุรกิจ (Business Rules)")
body("ทีมพัฒนาต้องการคำตอบสำหรับคำถามเหล่านี้ก่อนเริ่มพัฒนาแต่ละโมดูล:")

heading3("ด้านการซ่อมบำรุง")
bullet("ระยะกิโลเมตรที่ต้องแจ้งเตือนการซ่อมตามระยะคือเท่าไหร่? (ค่าเริ่มต้น: 80% ของรอบบำรุง)")
bullet("รอบบำรุงมาตรฐานของรถแต่ละประเภทคือกี่กิโลเมตร?")
bullet("ใครเป็นผู้อนุมัติการซ่อม และมีขั้นตอนอนุมัติกี่ระดับ?")

heading3("ด้านใบขับขี่และประกัน")
bullet("แจ้งเตือนล่วงหน้ากี่วันก่อนหมดอายุ? (ค่าเริ่มต้น: 30 วัน)")
bullet("มีประเภทประกันภัยอะไรบ้างที่ต้องติดตาม?")
bullet("ใบขับขี่ประเภทใดบ้างที่รถบัสต้องใช้?")

heading3("ด้านเชื้อเพลิง")
bullet("การคำนวณต้นทุนน้ำมันแยกตามบริษัทรถ หรือแยกตามโรงงาน หรือทั้งสองอย่าง?")
bullet("หน่วยราคาน้ำมันคิดอย่างไร? (ลิตร / บาท / คัน)")

heading3("ด้านการอบรม")
bullet("หลักสูตรใดบ้างที่บังคับ และมีวันหมดอายุ?")
bullet("ใครรับผิดชอบบันทึกผลการอบรม?")

heading3("ด้าน Notification")
bullet("ต้องการแจ้งเตือนผ่านช่องทางใดนอกจาก In-app? (เช่น Email / LINE)")
bullet("ใครได้รับการแจ้งเตือน: ผู้ดูแลระบบเท่านั้น หรือ operation ด้วย?")
bullet("เรื่องร้องเรียนที่ยังไม่ได้แก้ไข แจ้งเตือนหลังกี่วัน? (ค่าเริ่มต้น: 7 วัน)")

heading2("2.6 การอนุมัติและ Feedback ระหว่างการพัฒนา")
add_table(
    ["กิจกรรม", "รายละเอียด", "ผู้มีอำนาจอนุมัติ"],
    [
        ("อนุมัติ Wireframe / UI Design", "ต้องอนุมัติก่อนเริ่ม Code ในแต่ละ Phase", ""),
        ("อนุมัติ Database Schema", "ก่อนสร้าง Table จริงใน Supabase", ""),
        ("ทดสอบ UAT (User Acceptance Test)", "ลูกค้าทดสอบระบบจริงก่อน Go-Live ทุก Phase", ""),
        ("อนุมัติ Go-Live", "ยืนยันว่าพร้อม Deploy สู่ Production", ""),
        ("ตอบกลับ Bug Report", "ลูกค้าต้องยืนยันว่า Bug ได้รับการแก้ไขแล้ว", ""),
    ],
    col_widths=[5, 8, 3.5],
)
note_box("กรุณาระบุชื่อผู้มีอำนาจอนุมัติในแต่ละกิจกรรมก่อนเริ่มโครงการ")


# ═══════════════════════════════════════════════════════════
# 3. Functional Requirements
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1("3. ความต้องการเชิงฟังก์ชัน (Functional Requirements)")

modules = [
    ("FR-01", "Overview Dashboard",
     ["แสดง KPI รวม: จำนวนรถ, คนขับ, ต้นทุน, การแจ้งเตือนที่ค้างอยู่",
      "กราฟต้นทุนรายเดือน แยกรถแอร์/พัดลม",
      "สรุปสถานะแต่ละโมดูลในหน้าเดียว"]),
    ("FR-02", "Driver Management",
     ["บันทึกข้อมูลส่วนตัวพนักงานขับรถ",
      "ติดตามสถานะใบขับขี่และวันหมดอายุ",
      "บันทึกการตรวจสุขภาพและประวัติการอบรม",
      "กำหนดเส้นทาง/กะการทำงาน"]),
    ("FR-03", "Bus Fleet Management",
     ["ทะเบียนรถบัสทั้ง 68 คัน",
      "ติดตามสถานะรถ (ปกติ / ซ่อม / ปลดระวาง)",
      "บันทึกวันครบกำหนดประกันภัย",
      "ลิงก์ไปยังข้อมูลพฤติกรรมการขับขี่"]),
    ("FR-04", "Maintenance",
     ["บันทึกระยะทางรายวัน",
      "ตารางซ่อมบำรุงตามกำหนด",
      "บันทึกค่าซ่อม + แนบรูปถ่าย",
      "บันทึกค่าซ่อมจากอุบัติเหตุแยกต่างหาก"]),
    ("FR-05", "Cost Control",
     ["รวบรวมค่าใช้จ่ายทุกประเภท",
      "แยกแสดงต้นทุนรถแอร์ vs รถพัดลม",
      "Dashboard สรุปรายเดือน/ปี"]),
    ("FR-06", "Training Management",
     ["จัดการหลักสูตรอบรม",
      "บันทึกประวัติการอบรมรายคน",
      "แจ้งเตือนเมื่อใกล้วันหมดอายุ"]),
    ("FR-07", "License & Insurance",
     ["ติดตามวันหมดอายุใบขับขี่ / ประกันภัย",
      "บันทึกเคลมประกัน",
      "แนบเอกสาร (jpg/png/pdf ≤ 10MB)",
      "วิเคราะห์ต้นทุนการต่ออายุ"]),
    ("FR-08", "Customer Complaint",
     ["บันทึกเรื่องร้องเรียนพร้อมวันที่",
      "ติดตามสถานะการแก้ไข",
      "แจ้งเตือนเรื่องที่ค้างเกิน 7 วัน"]),
    ("FR-09", "Safety Measurement",
     ["Pre-trip Checklist",
      "บันทึกผลเป่าแอลกอฮอล์",
      "บันทึก Near-miss / เหตุการณ์เฉียดอุบัติเหตุ",
      "รายงานการขับรถเกินความเร็ว (Mock Data)"]),
    ("FR-10", "Purchase",
     ["รายการสั่งซื้ออุปกรณ์สำหรับรถแอร์",
      "รายการสั่งซื้อสำหรับสถานีเชื้อเพลิง"]),
    ("FR-11", "Fuel Management (ATG)",
     ["บันทึกการเติมน้ำมันรายวัน",
      "แยกตาม 7 บริษัทรถ",
      "รองรับการเชื่อมต่อระบบ ATG ในอนาคต"]),
    ("FR-12", "User Management (Admin Only)",
     ["CRUD ผู้ใช้งาน",
      "กำหนดบทบาท (admin / operation)",
      "Audit Log: บันทึกการสร้าง/แก้ไข/ลบข้อมูล"]),
]

for fr_id, name, reqs in modules:
    heading2(f"{fr_id} — {name}")
    for r in reqs:
        bullet(r)

doc.add_paragraph()
heading2("ฟีเจอร์ทั่วไปที่มีในทุกโมดูล")
bullets = [
    "Export ข้อมูลเป็น Excel (.xlsx) และ PDF",
    "ค้นหาและกรองข้อมูล",
    "Pagination สำหรับข้อมูลจำนวนมาก",
    "Responsive บนมือถือ (ดูข้อมูลได้ ไม่จำเป็นต้องแก้ไขบนมือถือ)",
    "ภาษาไทยทั้งหมด (UI, ข้อความ, การแจ้งเตือน)",
]
for b in bullets:
    bullet(b)


# ═══════════════════════════════════════════════════════════
# 4. Non-Functional Requirements
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1("4. ความต้องการที่ไม่ใช่เชิงฟังก์ชัน (Non-Functional Requirements)")

add_table(
    ["หัวข้อ", "ข้อกำหนด"],
    [
        ("ประสิทธิภาพ (Performance)", "หน้าโหลดภายใน 3 วินาทีในสภาพเครือข่ายปกติ"),
        ("ความพร้อมใช้งาน (Availability)", "Uptime ≥ 99% (Vercel + Supabase SLA)"),
        ("ความปลอดภัย (Security)", "HTTPS ทุก endpoint, Row Level Security (RLS) บน Supabase, ไม่เก็บรหัสผ่านในระบบ"),
        ("ขนาดไฟล์แนบ", "สูงสุด 10 MB ต่อไฟล์ รองรับ jpg, png, pdf"),
        ("Browser Support", "Chrome, Edge, Firefox เวอร์ชันล่าสุด 2 รุ่น"),
        ("อุปกรณ์หลัก", "PC / Laptop (1920×1080) — Mobile เป็น Responsive"),
        ("ภาษา", "Thai UI ทั้งหมด"),
        ("Backup", "Supabase จัดการ Backup อัตโนมัติ (Point-in-time Recovery)"),
        ("Audit Trail", "บันทึก Create/Update/Delete พร้อม Before/After values สำหรับ Admin"),
    ],
    col_widths=[6, 10.5],
)


# ═══════════════════════════════════════════════════════════
# 5. สถาปัตยกรรมระบบ
# ═══════════════════════════════════════════════════════════
heading1("5. สถาปัตยกรรมระบบ (System Architecture)")

add_table(
    ["Layer", "เทคโนโลยี", "บทบาท"],
    [
        ("Frontend", "React + Tailwind CSS", "UI ทั้งหมด, ภาษาไทย"),
        ("Backend / Database", "Supabase (PostgreSQL)", "เก็บข้อมูล, Auth, Storage, Realtime"),
        ("Authentication", "Supabase Auth (Email/Password)", "Login + Role-based Access"),
        ("File Storage", "Supabase Storage", "ไฟล์แนบ (รูปภาพ, PDF)"),
        ("Deployment", "Vercel", "Host Frontend, CI/CD อัตโนมัติ"),
        ("Version Control", "GitHub", "Source code management"),
    ],
    col_widths=[3.5, 5, 8],
)

note_box("โครงสร้างนี้ใช้บริการ Cloud ทั้งหมด — ลูกค้าไม่ต้องจัดหา Server เอง")


# ═══════════════════════════════════════════════════════════
# 6. Timeline
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1("6. แผนการดำเนินงาน (Proposed Timeline)")

body("Timeline ด้านล่างเป็นแผนเบื้องต้น ระยะเวลาจริงจะถูกตกลงร่วมกันหลังได้รับข้อมูลครบถ้วน")

add_table(
    ["Phase", "รายละเอียด", "โมดูล", "ระยะเวลา (โดยประมาณ)"],
    [
        ("Phase 1", "Project Setup + Auth", "User Management, Login, Dashboard (โครง)", "2 สัปดาห์"),
        ("Phase 2", "Driver Module", "Driver Management", "2 สัปดาห์"),
        ("Phase 3", "Bus Module", "Bus Fleet Management", "1 สัปดาห์"),
        ("Phase 4", "Maintenance Module", "Maintenance + Cost Control", "3 สัปดาห์"),
        ("Phase 5", "Compliance & Safety", "License & Insurance, Training, Safety", "3 สัปดาห์"),
        ("Phase 6", "Fuel & Purchase", "Fuel (ATG), Purchase, Complaint", "2 สัปดาห์"),
        ("Phase 7", "Dashboard & Reports", "Overview Dashboard, Export ทุกโมดูล", "2 สัปดาห์"),
        ("Phase 8", "UAT & Go-Live", "ทดสอบทั้งระบบ, แก้ไข Bug, Deploy", "2 สัปดาห์"),
    ],
    col_widths=[2, 4.5, 6, 4],
)
note_box("ลูกค้าต้องพร้อมทำ UAT ภายในระยะเวลาที่ตกลงกัน เพื่อไม่ให้ Timeline ล่าช้า")


# ═══════════════════════════════════════════════════════════
# 7. Checklist สรุปสำหรับลูกค้า
# ═══════════════════════════════════════════════════════════
heading1("7. Checklist สรุป — สิ่งที่ลูกค้าต้องส่งให้ทีมพัฒนา")

body("☐ = ยังไม่ได้ส่ง   ☑ = ส่งแล้ว")
doc.add_paragraph()

checklist_items = [
    ("ข้อมูลองค์กร", [
        "ชื่อบริษัท (ไทย + อังกฤษ)",
        "โลโก้บริษัท (PNG/SVG ≥ 512px)",
        "รายชื่อและที่อยู่ 3 โรงงาน",
        "รายชื่อ 7 บริษัทรถบัส",
    ]),
    ("ข้อมูลผู้ใช้งาน", [
        "รายชื่อผู้ใช้งานทั้งหมด (ชื่อ + อีเมล + บทบาท)",
        "อีเมล Admin คนแรก",
        "ระบุผู้มีอำนาจอนุมัติในแต่ละขั้นตอน",
    ]),
    ("Master Data (Excel)", [
        "ข้อมูลพนักงานขับรถทั้งหมด",
        "ข้อมูลรถบัส 68 คัน (ทะเบียน, ยี่ห้อ, รุ่น, ประเภท, โรงงาน)",
        "รายชื่อหลักสูตรอบรม",
        "ข้อมูลประวัติที่ต้องการ migrate (ถ้ามี)",
    ]),
    ("เอกสารและแบบฟอร์ม", [
        "แบบฟอร์มตรวจสภาพก่อนออกรถ (Safety Checklist)",
        "แบบฟอร์มบันทึกอุบัติเหตุ",
        "ตัวอย่างเอกสารประกันภัย",
    ]),
    ("Business Rules", [
        "ระยะกิโลเมตรแจ้งเตือนการซ่อม",
        "จำนวนวันแจ้งเตือนใบขับขี่/ประกันหมดอายุ",
        "ขั้นตอนการอนุมัติซ่อมบำรุง",
        "ช่องทาง Notification ที่ต้องการ",
    ]),
    ("เทคนิค", [
        "โดเมนเนม (ถ้าต้องการใช้โดเมนเอง)",
        "ข้อมูล ATG API (สำหรับ Phase อนาคต)",
        "ข้อมูล Driving Behavior System (สำหรับ Phase อนาคต)",
    ]),
]

for category, items in checklist_items:
    heading3(f"□ {category}")
    for item in items:
        bullet(f"☐  {item}")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════
# 8. ข้อตกลงและลายเซ็น
# ═══════════════════════════════════════════════════════════
doc.add_page_break()
heading1("8. การรับรองเอกสาร")

body("เอกสาร SRS ฉบับนี้ได้รับการตรวจสอบและอนุมัติโดยผู้มีอำนาจจากทั้งสองฝ่าย")

doc.add_paragraph()

sig_table = doc.add_table(rows=6, cols=2)
sig_table.style = "Table Grid"
sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sig_data = [
    ("ฝ่ายลูกค้า (ผู้ว่าจ้าง)", "ฝ่ายพัฒนา (ผู้รับจ้าง)"),
    ("ชื่อ: ___________________________", "ชื่อ: ___________________________"),
    ("ตำแหน่ง: ________________________", "ตำแหน่ง: ________________________"),
    ("ลายเซ็น: ________________________", "ลายเซ็น: ________________________"),
    ("วันที่: __ / __ / ____", "วันที่: __ / __ / ____"),
    ("ประทับตรา (ถ้ามี)", "ประทับตรา (ถ้ามี)"),
]

for ri, (left, right) in enumerate(sig_data):
    for ci, text in enumerate([left, right]):
        cell = sig_table.rows[ri].cells[ci]
        if ri == 0:
            shade_cell(cell, "1E3A8A")
            set_cell_text(cell, text, bold=True, color=WHITE, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
        else:
            set_cell_text(cell, text, size=13)
        cell.width = Cm(8)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— จบเอกสาร —")
set_font(run, 13, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("เอกสารนี้จัดทำโดยทีมพัฒนาระบบ BMS  |  v1.0  |  " + datetime.date.today().strftime("%d/%m/%Y"))
set_font(run, 11, color=GRAY)

# ─── Save ───
out_path = r"C:\Users\nakrit\Desktop\BMS\BMS_SRS_Client_Requirements_v1.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
